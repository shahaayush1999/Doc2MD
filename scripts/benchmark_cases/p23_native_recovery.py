from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader, PdfWriter

from .common import (
    REPO_ROOT,
    directed_edge_leaf,
    finalize_fact_regions,
    form_state_leaf,
    leaf,
    markdown_table,
    table_leaves,
    visual_leaf,
)


CASE_ID = "P23-native-text-layer-recovery"
TITLE = "Northstar Cold Chain Supplier Cutover Authorization"
FAMILY = "office export recovery"
TAGS = [
    "native-pdf",
    "libreoffice-export",
    "malformed-layout",
    "cross-office-conversion",
    "overlapping-text-boxes",
    "font-substitution",
    "detached-labels",
    "scrambled-reading-order",
    "broken-object-stacking",
    "native-text-recovery",
    "mixed-modality",
    "tables",
    "source-precedence",
    "reading-order",
]


CONTROL_HEADERS = ["Field", "Controlled value", "Field", "Controlled value"]
CONTROL_ROWS = [
    ["File", "SV-2048", "Revision", "E"],
    ["Issued", "12 Jul 2026 18:40 MST", "Cutover start", "14 Jul 2026 22:00 MST"],
    ["Controlled lots", "PX-771 and TZ-219", "Coordinator", "Avery Kim"],
]

AUTHORITY_HEADERS = ["Priority", "Record", "Controls", "Does not control"]
AUTHORITY_ROWS = [
    ["1", "Page 5 executed authorization", "Selective GO/HOLD, exclusions, signed corrections", "Rates or exception closure"],
    ["2", "Page 4 dependency / exception sheet", "Directed gates, exception state, accepted evidence", "Commercial ceiling"],
    ["3", "Page 3 approved reconciliation", "Current rates, ceiling, and invoice route", "Operational release"],
    ["4", "Page 2 cutover workplan", "Sequence, windows, handoffs, and task gates", "Final authorization"],
]

WORKPLAN_HEADERS = ["ID", "Group", "Site / asset", "Activity", "Owner", "Window", "Gate / evidence", "State"]
WORKPLAN_ROWS = [
    [
        "T-01",
        "PHOENIX",
        "Phoenix / Dock 4",
        "Reserve the controlled bay and verify that Kestrel access list AC-44 remains active.",
        "Avery Kim",
        "14 Jul 21:15-21:45",
        "FC-119 acknowledged; AC-44 posted",
        "READY",
    ],
    [
        "T-02",
        "PHX (cont.)",
        "Phoenix / Sensor P-14",
        "Shift live telemetry from KG-2 to BG-5 while maintaining parallel capture.",
        "Lin Chen",
        "14 Jul 21:45-22:25",
        "TV-19 delta <=0.3 C for 30 continuous min",
        "CONDITIONAL",
    ],
    [
        "T-04",
        "PHX (cont.)",
        "Phoenix / Lot PX-771",
        "Transfer 48 eligible pallets after T-02 and condition C-1 release the lot.",
        "Rina Singh",
        "14 Jul 22:30-00:10",
        "SC-77: 48 moved; 0 orphan pallets",
        "BLOCKED: T-02",
    ],
    [
        "T-03",
        "TUCSON",
        "Tucson / Cage C",
        "Seal Kestrel returns and complete a dual count before custody release.",
        "Mateo Soto",
        "14 Jul 22:05-22:45",
        "IC-882 signed by Northstar and carrier",
        "READY",
    ],
    [
        "T-05",
        "TUS (cont.)",
        "Tucson / Lot TZ-219",
        "Transfer 26 eligible pallets; keep six QA-61 cases segregated.",
        "Jordan Cole",
        "14 Jul 22:50-00:00",
        "QR-203 or segregated-custody receipt",
        "READY: EXCLUDE 6",
    ],
    [
        "T-06",
        "BOTH SITES",
        "Both sites / credentials",
        "Reconcile both lots, then close Kestrel credentials only after the waiver.",
        "Avery Kim",
        "15 Jul 00:20-01:00",
        "REC-2048 zero unresolved + EX-07 signed",
        "HOLD",
    ],
]

HANDOFF_HEADERS = ["Trigger", "Notify", "Channel", "Required content"]
HANDOFF_ROWS = [
    ["T-02 passes", "Phoenix floor lead", "BR-2048", "P-14 final delta and validation end time"],
    ["T-04 completes", "Inventory control", "IC-WEST", "PX-771 moved count and orphan count"],
    ["RB-12 invoked", "Both site leads", "BR-2048", "Last completed task and custody location"],
]

RECON_HEADERS = ["Line", "Submitted / superseded", "Approved / current", "Approval note"]
RECON_ROWS = [
    ["Boreal implementation", "Q-884 rev B / $19,200", "Q-884 rev C / $18,600", "Fixed fee"],
    ["Kestrel overlap", "3 nights / $6,480", "Maximum 2 nights / $4,320", "Only while EX-07 remains open"],
    ["Sensor validation", "TV-19 / $2,150", "TV-19 / $2,150", "No change"],
    ["Contingency", "$4,500", "$3,000", "Only if EX-07 extends overlap"],
    ["Request / ceiling", "Submitted total / $32,330", "NTE ceiling / $28,070", "Reapproval above ceiling"],
]

INVOICE_HEADERS = ["Vendor", "PO", "Cost center", "Reviewer", "Route"]
INVOICE_ROWS = [
    ["Boreal Storage", "44018-3", "CC-7420", "Jonah Mercer", "AP-West / transition"],
    ["Kestrel Logistics", "44007-9", "CC-7420", "Jonah Mercer", "AP-West / overlap"],
    ["ThermoVision", "43991-2", "CC-7314", "Nadia Brooks", "AP-Quality / validation"],
]

EXCEPTION_HEADERS = [
    "Exception",
    "Affected obligation",
    "Current finding",
    "Owner",
    "Due",
    "Accepted closure evidence",
    "State",
]
EXCEPTION_ROWS = [
    [
        "E-05",
        "T-02 telemetry gate",
        "Parallel baseline differs by 0.2 C after probe alignment.",
        "Lin Chen",
        "12 Jul 15:00 MST",
        "Signed TV-19 trace with 30-minute comparison",
        "CLOSED",
    ],
    [
        "EX-07",
        "T-06 credential closure",
        "Kestrel access waiver is with counsel; physical cutover may proceed while access is retained.",
        "Isha Patel",
        "14 Jul 18:00 MST",
        "Countersigned waiver WA-771",
        "OPEN",
    ],
    [
        "E-08",
        "T-05 Tucson inventory",
        "Six QA-61 cases remain quarantined and are outside the eligible transfer quantity.",
        "Dara Nwosu",
        "14 Jul 22:45 MST",
        "QR-203 release or segregated custody receipt",
        "OPEN - EXCLUDE",
    ],
    [
        "E-09",
        "Boreal invoice route",
        "Draft purchase order omitted the transition suffix.",
        "Jonah Mercer",
        "09 Jul 12:00 MST",
        "Issued PO 44018-3",
        "CLOSED",
    ],
    [
        "E-11",
        "T-03 carrier release",
        "Carrier arrival is forecast at 22:55, ten minutes after the planned count window.",
        "Mateo Soto",
        "14 Jul 22:15 MST",
        "Dispatch update in BR-2048 and revised custody time",
        "WATCH",
    ],
    ["E-14", "T-04 scan evidence", "Scanner SC-77B is the approved fallback if SC-77 loses sync.", "Rina Singh", "14 Jul 22:20 MST", "Both batch IDs recorded in IC-WEST", "MONITOR"],
]

GRAPH_EDGES = [
    ("T-01", "T-02", "bay ready"),
    ("T-02", "C-1", "delta proof"),
    ("C-1", "T-04", "releases"),
    ("T-02", "T-04", "gates"),
    ("T-03", "T-05", "custody"),
    ("E-08", "T-05", "exclude 6"),
    ("C-2", "T-05", "eligible only"),
    ("T-04", "T-06", "reconcile PX"),
    ("T-05", "T-06", "reconcile TZ"),
    ("EX-07", "T-06", "waiver"),
    ("C-3", "T-06", "close access"),
    ("C-4", "RB-12", "breach"),
    ("RB-12", "T-04", "halts"),
    ("RB-12", "T-05", "halts"),
]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 50, start: int = 60, bottom: int = 45, end: int = 60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_fixed(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "3")
        node.set(qn("w:color"), "B8C0C7")
        borders.append(node)


def _write_cell(cell, value: str, *, header: bool = False, compact: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell, top=28 if compact else 45, bottom=24 if compact else 40)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(7.8 if compact else 9.3)
    run = paragraph.add_run(str(value))
    run.bold = header
    run.font.name = "Arial"
    run.font.size = Pt(6.7 if compact else 7.5)
    run.font.color.rgb = RGBColor(25, 35, 43)


def _add_edge_label_fragments(document: Document) -> None:
    """Add only the native connector-ID-to-label fragments.

    Endpoints and direction exist only in the visible graph.  The text layer
    carries the detached label objects, each keyed by the connector ID printed
    on the arrow.  Neither modality is a complete answer key by itself.
    """
    objects = [
        {
            "ordinal": index - 1,
            "text": f"E{index:02d}  {relation}",
            # A damaged slide import has collapsed the connector-label band:
            # neighboring opaque boxes overlap, so page pixels reveal only
            # fragments while the PDF's native object stream retains each
            # complete ID/label pair.
            "x": 0.48 + ((index - 1) % 7) * 0.36,
            "y": 5.02 + ((index - 1) // 7) * 0.16,
            "font_size": 23.0 if index in {3, 6, 10, 13} else 18.0,
            "fill": "F4E8CE" if index in {6, 7, 10, 11} else "E7EEF2",
        }
        for index, (_, _, relation) in enumerate(GRAPH_EDGES, start=1)
    ]
    # Short labels would otherwise survive intact at the top of the collapsed
    # band.  Import/z-order damage places the longer E11 object above E10/E12/E13,
    # leaving every one as a visible fragment while retaining all complete
    # native objects.
    insertion_order = [objects[9], objects[12], objects[11], objects[10], objects[13]] + objects[:9]
    for object_index, item in enumerate(insertion_order):
        _add_floating_cell(
            document,
            item["text"],
            x_inches=item["x"],
            y_inches=item["y"],
            width_inches=0.86,
            height_inches=0.14,
            font_size=item["font_size"],
            font_name="Aptos Narrow Variable" if item["font_size"] > 9 else "Arial",
            fill=item["fill"],
            bold=True,
        )
        if object_index == 3:
            _add_floating_panel(
                document,
                "REV E / CONTROL PATH LABEL BAND",
                x_twips=2050,
                y_twips=7160,
                width_inches=5.9,
                height_inches=0.44,
                font_size=20,
                fill="C9D9E1",
                color=RGBColor(44, 79, 96),
            )


def _keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)


def _label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2.5)
    _keep_with_next(paragraph)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(6.4)
    run.font.color.rgb = RGBColor(68, 82, 93)


def _paragraph(document: Document, text: str, *, size: float = 8.4, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = Pt(size + 2.1)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(25, 35, 43)


def _page_title(document: Document, page: int, title: str, status: str) -> None:
    banner = document.add_table(rows=1, cols=2)
    banner.autofit = False
    banner.alignment = WD_TABLE_ALIGNMENT.LEFT
    banner.columns[0].width = Inches(6.15 if page != 4 else 8.7)
    banner.columns[1].width = Inches(1.25)
    for cell in banner.rows[0].cells:
        _set_cell_margins(cell, top=5, bottom=5, start=0, end=0)
    left = banner.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("NORTHSTAR COLD CHAIN  |  SUPPLIER TRANSITION SV-2048")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(6.4)
    run.font.color.rgb = RGBColor(60, 76, 87)
    right = banner.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run(f"{status}  |  {page}/5")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(6.4)
    run.font.color.rgb = RGBColor(124, 55, 42)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    if page <= 3:
        run.font.name = "Aptos Display Condensed"
        run.font.size = Pt({1: 35, 2: 27, 3: 31}[page])
        paragraph.paragraph_format.line_spacing = Pt(11)
        paragraph.paragraph_format.space_after = Pt(0)
    else:
        run.font.name = "Arial"
        run.font.size = Pt(16.5)
    run.font.color.rgb = RGBColor(19, 38, 52)
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(2)
    run = rule.add_run("_" * (96 if page == 4 else 72))
    run.font.name = "Arial"
    run.font.size = Pt(5.6)
    run.font.color.rgb = RGBColor(170, 181, 188)


def _configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.38)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.42)
        section.right_margin = Inches(0.42)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.42)
        section.bottom_margin = Inches(0.42)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)



def _add_floating_panel(
    document: Document,
    text: str,
    *,
    x_twips: int,
    y_twips: int,
    width_inches: float,
    height_inches: float,
    font_size: float,
    fill: str = "F4E8CE",
    color: RGBColor = RGBColor(118, 65, 42),
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Add a page-anchored table that behaves like an imported slide text box.

    The use of floating tables is deliberate: Word and Writer disagree on their
    wrap geometry and object order, closely reproducing the overlapping boxes
    seen in damaged Slides/PowerPoint round trips while preserving native text.
    """
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(width_inches)
    _set_table_fixed(table)
    tbl_pr = table._tbl.tblPr
    position = OxmlElement("w:tblpPr")
    position.set(qn("w:leftFromText"), "0")
    position.set(qn("w:rightFromText"), "0")
    position.set(qn("w:topFromText"), "0")
    position.set(qn("w:bottomFromText"), "0")
    position.set(qn("w:vertAnchor"), "page")
    position.set(qn("w:horzAnchor"), "page")
    position.set(qn("w:tblpX"), str(x_twips))
    position.set(qn("w:tblpY"), str(y_twips))
    tbl_pr.append(position)
    overlap = OxmlElement("w:tblOverlap")
    overlap.set(qn("w:val"), "overlap")
    tbl_pr.append(overlap)
    row = table.rows[0]
    row.height = Inches(height_inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Inches(width_inches)
    _set_cell_shading(cell, fill)
    _write_cell(cell, text, header=True, compact=True)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.runs[0]
    run.font.name = "Aptos Display Condensed"
    run.font.size = Pt(font_size)
    run.font.color.rgb = color


def _add_floating_cell(
    document: Document,
    text: str,
    *,
    x_inches: float,
    y_inches: float,
    width_inches: float,
    height_inches: float,
    font_size: float = 7.3,
    font_name: str = "Arial",
    fill: str = "FFFFFF",
    color: RGBColor = RGBColor(25, 35, 43),
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Place one native cell as an independent page-anchored object.

    A genuine damaged slide/deck export does not necessarily retain a logical
    table object: every cell may survive as its own text box, and PDF extraction
    then follows object/z-order rather than row order.  This helper intentionally
    models that failure while leaving the glyphs and their page coordinates
    recoverable.  It is fixture construction only; the runtime does not know
    about this representation.
    """
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(width_inches)
    _set_table_fixed(table)
    tbl_pr = table._tbl.tblPr
    position = OxmlElement("w:tblpPr")
    position.set(qn("w:leftFromText"), "0")
    position.set(qn("w:rightFromText"), "0")
    position.set(qn("w:topFromText"), "0")
    position.set(qn("w:bottomFromText"), "0")
    position.set(qn("w:vertAnchor"), "page")
    position.set(qn("w:horzAnchor"), "page")
    position.set(qn("w:tblpX"), str(int(x_inches * 1440)))
    position.set(qn("w:tblpY"), str(int(y_inches * 1440)))
    tbl_pr.append(position)
    overlap = OxmlElement("w:tblOverlap")
    overlap.set(qn("w:val"), "overlap")
    tbl_pr.append(overlap)

    row = table.rows[0]
    row.height = Inches(height_inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Inches(width_inches)
    _set_cell_shading(cell, fill)
    _write_cell(cell, text, compact=True)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = Pt(max(2.3, min(font_size + 0.8, 9.0)))
    run = paragraph.runs[0]
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = color


def _add_spatial_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    x: float,
    y: float,
    header_height: float,
    row_height: float,
    seed: int,
    body_height: float | None = None,
    omitted_headers: set[int] | None = None,
    damage_sizes: dict[tuple[int, int], float] | None = None,
    cell_offsets: dict[tuple[int, int], tuple[float, float]] | None = None,
    bold_columns: set[int] | None = None,
) -> None:
    """Draw a table as independent native objects in scrambled object order.

    Coordinates retain the intended row/column relationships.  The PDF text
    stream deliberately does not: it is the same failure mode produced when a
    presentation table is decomposed into text boxes during an office round
    trip.  No value is hidden or moved outside the page.
    """
    omitted_headers = omitted_headers or set()
    damage_sizes = damage_sizes or {}
    cell_offsets = cell_offsets or {}
    bold_columns = bold_columns or set()
    body_height = body_height or row_height
    x_positions: list[float] = []
    cursor = x
    for width in widths:
        x_positions.append(cursor)
        cursor += width

    objects: list[dict] = []
    ordinal = 0
    for column_index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        objects.append(
            {
                "ordinal": ordinal,
                "row": -1,
                "column": column_index,
                "text": "" if column_index in omitted_headers else header,
                "x": x_positions[column_index],
                "y": y,
                "width": width,
                "height": header_height,
                "font_size": 6.8,
                "font_name": "Arial",
                "fill": "C8D8E1",
                "bold": True,
            }
        )
        ordinal += 1

    for row_index, values in enumerate(rows):
        for column_index, (value, width) in enumerate(zip(values, widths, strict=True)):
            x_offset, y_offset = cell_offsets.get((row_index, column_index), (0.0, 0.0))
            font_size = damage_sizes.get((row_index, column_index), 7.0)
            objects.append(
                {
                    "ordinal": ordinal,
                    "row": row_index,
                    "column": column_index,
                    "text": str(value),
                    "x": x_positions[column_index] + x_offset,
                    "y": y + header_height + row_index * row_height + y_offset,
                    "width": width,
                    "height": body_height,
                    "font_size": font_size,
                    "font_name": "Aptos Narrow Variable" if font_size > 8.2 else "Arial",
                    "fill": "F2F5F6" if row_index % 2 else "FFFFFF",
                    "bold": column_index in bold_columns,
                }
            )
            ordinal += 1

    column_order = list(range(len(headers)))
    column_shift = seed % len(column_order)
    column_order = column_order[column_shift:] + column_order[:column_shift]
    if seed % 2:
        column_order.reverse()
    row_order = list(range(-1, len(rows)))
    row_shift = (seed // 7) % len(row_order)
    row_order = row_order[row_shift:] + row_order[:row_shift]
    if (seed // 3) % 2:
        row_order.reverse()
    object_by_position = {(int(item["row"]), int(item["column"])): item for item in objects}
    insertion_order = [object_by_position[(row_index, column_index)] for column_index in column_order for row_index in row_order]

    for item in insertion_order:
        _add_floating_cell(
            document,
            item["text"],
            x_inches=item["x"],
            y_inches=item["y"],
            width_inches=item["width"],
            height_inches=item["height"],
            font_size=item["font_size"],
            font_name=item["font_name"],
            fill=item["fill"],
            bold=item["bold"],
        )


def _add_floating_picture(
    document: Document,
    path: Path,
    *,
    x_twips: int,
    y_twips: int,
    width_inches: float,
) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(width_inches)
    tbl_pr = table._tbl.tblPr
    position = OxmlElement("w:tblpPr")
    position.set(qn("w:leftFromText"), "0")
    position.set(qn("w:rightFromText"), "0")
    position.set(qn("w:topFromText"), "0")
    position.set(qn("w:bottomFromText"), "0")
    position.set(qn("w:vertAnchor"), "page")
    position.set(qn("w:horzAnchor"), "page")
    position.set(qn("w:tblpX"), str(x_twips))
    position.set(qn("w:tblpY"), str(y_twips))
    tbl_pr.append(position)
    overlap = OxmlElement("w:tblOverlap")
    overlap.set(qn("w:val"), "overlap")
    tbl_pr.append(overlap)
    cell = table.cell(0, 0)
    _set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))


def _font(size: int, *, bold: bool = False, italic: bool = False):
    if bold and italic:
        arial_name = "Arial Bold Italic.ttf"
        verdana_name = "Verdana Bold Italic.ttf"
    elif bold:
        arial_name = "Arial Bold.ttf"
        verdana_name = "Verdana Bold.ttf"
    elif italic:
        arial_name = "Arial Italic.ttf"
        verdana_name = "Verdana Italic.ttf"
    else:
        arial_name = "Arial.ttf"
        verdana_name = "Verdana.ttf"
    candidates = [
        f"/System/Library/Fonts/Supplemental/{arial_name}",
        f"/System/Library/Fonts/Supplemental/{verdana_name}",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default(size=size)


def _finance_stamp(path: Path) -> None:
    image = Image.new("RGBA", (920, 300), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    green = (36, 104, 76, 238)
    draw.rounded_rectangle((15, 15, 905, 285), radius=24, outline=green, width=9)
    draw.line((42, 96, 878, 96), fill=green, width=4)
    draw.text((46, 36), "FINANCE VALIDATED", font=_font(42, bold=True), fill=green)
    draw.text((46, 120), "JONAH MERCER  |  09 JUL 2026 16:42 MST", font=_font(28), fill=green)
    draw.text((46, 171), "APPROVED RECONCILIATION  |  CONTROL RC-2048-E", font=_font(28, bold=True), fill=green)
    draw.text((46, 222), "COST VALIDATION IS NOT OPERATIONAL RELEASE", font=_font(24), fill=green)
    image.save(path, format="PNG", optimize=True)



def _arrow_path(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[float, float]],
    *,
    color: tuple[int, int, int],
    width: int = 4,
) -> None:
    if len(points) < 2:
        return
    adjusted = [list(point) for point in points]
    sx, sy = adjusted[0]
    nx, ny = adjusted[1]
    length = max(((nx - sx) ** 2 + (ny - sy) ** 2) ** 0.5, 1)
    adjusted[0] = [sx + (nx - sx) * 62 / length, sy + (ny - sy) * 62 / length]
    px, py = adjusted[-2]
    ex, ey = adjusted[-1]
    length = max(((ex - px) ** 2 + (ey - py) ** 2) ** 0.5, 1)
    end = [ex - (ex - px) * 62 / length, ey - (ey - py) * 62 / length]
    adjusted[-1] = end
    polyline = [(int(x), int(y)) for x, y in adjusted]
    draw.line(polyline, fill=color, width=width, joint="curve")
    px, py = adjusted[-2]
    ex, ey = adjusted[-1]
    angle = __import__("math").atan2(ey - py, ex - px)
    spread = 0.55
    size = 15
    left = (ex - size * __import__("math").cos(angle - spread), ey - size * __import__("math").sin(angle - spread))
    right = (ex - size * __import__("math").cos(angle + spread), ey - size * __import__("math").sin(angle + spread))
    draw.polygon([(int(ex), int(ey)), (int(left[0]), int(left[1])), (int(right[0]), int(right[1]))], fill=color)


def _dependency_graph(path: Path) -> None:
    image = Image.new("RGB", (1500, 700), (250, 252, 253))
    draw = ImageDraw.Draw(image)
    navy = (25, 52, 70)
    blue = (47, 104, 137)
    amber = (171, 103, 34)
    red = (158, 58, 52)
    green = (42, 112, 78)
    draw.text((30, 22), "DIRECTED RELEASE LOGIC", font=_font(28, bold=True), fill=navy)
    draw.text((420, 27), "Arrowheads and connector IDs define direction; labels survived as detached objects.", font=_font(19), fill=(75, 88, 97))
    for y, label in ((95, "PHOENIX"), (305, "TUCSON"), (525, "EXCEPTIONS / ROLLBACK")):
        draw.text((24, y), label, font=_font(16, bold=True), fill=(98, 110, 118))

    nodes = {
        "T-01": (135, 140),
        "T-02": (380, 140),
        "C-1": (630, 140),
        "T-04": (910, 140),
        "T-06": (1400, 285),
        "T-03": (135, 355),
        "E-08": (380, 355),
        "C-2": (630, 355),
        "T-05": (910, 355),
        "C-4": (135, 575),
        "RB-12": (380, 575),
        "EX-07": (910, 575),
        "C-3": (1160, 575),
    }
    routes: dict[tuple[str, str], list[tuple[float, float]]] = {
        ("T-02", "T-04"): [nodes["T-02"], (420, 65), (870, 65), nodes["T-04"]],
        ("T-03", "T-05"): [nodes["T-03"], (180, 275), (870, 275), nodes["T-05"]],
        ("E-08", "T-05"): [nodes["E-08"], (420, 440), (870, 440), nodes["T-05"]],
        ("RB-12", "T-04"): [nodes["RB-12"], (430, 500), (840, 500), (840, 185), nodes["T-04"]],
        ("RB-12", "T-05"): [nodes["RB-12"], (430, 645), (840, 645), (840, 400), nodes["T-05"]],
    }
    edge_id_positions = {
        ("T-01", "T-02"): (255, 122),
        ("T-02", "C-1"): (505, 122),
        ("C-1", "T-04"): (770, 122),
        ("T-02", "T-04"): (630, 43),
        ("T-03", "T-05"): (535, 255),
        ("E-08", "T-05"): (650, 420),
        ("C-2", "T-05"): (770, 337),
        ("T-04", "T-06"): (1115, 185),
        ("T-05", "T-06"): (1130, 330),
        ("EX-07", "T-06"): (1130, 485),
        ("C-3", "T-06"): (1305, 445),
        ("C-4", "RB-12"): (255, 557),
        ("RB-12", "T-04"): (790, 480),
        ("RB-12", "T-05"): (650, 625),
    }
    edge_font = _font(16, bold=True)
    for edge_index, (source, destination, relation) in enumerate(GRAPH_EDGES, start=1):
        color = red if relation in {"breach", "halts"} else amber if relation in {"exclude 6", "eligible only", "waiver", "close access"} else blue
        points = routes.get((source, destination), [nodes[source], nodes[destination]])
        _arrow_path(draw, points, color=color)

    node_font = _font(20, bold=True)
    for name, (x, y) in nodes.items():
        if name.startswith("T-"):
            stroke, fill = blue, (225, 239, 247)
        elif name.startswith("C-"):
            stroke, fill = green, (227, 243, 234)
        elif name == "RB-12":
            stroke, fill = red, (249, 228, 226)
        else:
            stroke, fill = amber, (250, 238, 218)
        draw.rounded_rectangle((x - 59, y - 28, x + 59, y + 28), radius=12, fill=fill, outline=stroke, width=4)
        draw.text((x, y), name, font=node_font, fill=stroke, anchor="mm")

    # Opaque imported object groups sit above the original arrow stack. The
    # raster keeps endpoint geometry, arrowheads, and connector IDs; detached
    # native E##/relation fragments on physical pages 5-6 supply the labels.
    overlays = [
        ((535, 180, 845, 285), "PHOENIX RELEASE\nGATES", (232, 226, 214), (96, 74, 51)),
        ((1030, 345, 1315, 465), "EXCEPTION\nOVERRIDES", (244, 220, 215), (148, 57, 49)),
        ((255, 485, 700, 545), "ROLLBACK CONTROL RB-12", (220, 232, 236), (44, 79, 96)),
    ]
    for (x0, y0, x1, y1), text, fill, ink in overlays:
        draw.rectangle((x0, y0, x1, y1), fill=fill, outline=ink, width=3)
        lines = text.split("\n")
        for line_index, line in enumerate(lines):
            draw.text((x0 + 18, y0 + 24 + line_index * 31), line, font=_font(21, bold=True), fill=ink)
    # The connector ID tags survived above the corrupted group stack even where
    # the middle of a path is covered.  This keeps every endpoint binding
    # recoverable without restoring the detached relation label itself.
    for edge_index, (source, destination, relation) in enumerate(GRAPH_EDGES, start=1):
        color = red if relation in {"breach", "halts"} else amber if relation in {"exclude 6", "eligible only", "waiver", "close access"} else blue
        connector_id = f"E{edge_index:02d}"
        tx, ty = edge_id_positions[(source, destination)]
        bounds = draw.textbbox((tx, ty), connector_id, font=edge_font, anchor="mm")
        draw.rounded_rectangle((bounds[0] - 5, bounds[1] - 2, bounds[2] + 5, bounds[3] + 2), radius=5, fill=(250, 252, 253))
        draw.text((tx, ty), connector_id, font=edge_font, fill=color, anchor="mm")
    image.save(path, format="PNG", optimize=True)


def _executed_authorization(path: Path) -> None:
    image = Image.new("RGB", (1600, 1900), (240, 241, 238))
    draw = ImageDraw.Draw(image)
    paper = (255, 255, 252)
    ink = (28, 39, 46)
    muted = (89, 99, 104)
    blue_ink = (35, 76, 134)
    green = (33, 101, 70)
    red = (153, 55, 48)
    draw.rectangle((38, 28, 1562, 1872), fill=paper, outline=(178, 181, 176), width=3)
    draw.text((76, 66), "NORTHSTAR COLD CHAIN", font=_font(27, bold=True), fill=ink)
    draw.text((76, 108), "EXECUTED SELECTIVE CUTOVER AUTHORIZATION", font=_font(38, bold=True), fill=ink)
    draw.text((76, 158), "SV-2048  |  REV E  |  EXECUTED COPY 01", font=_font(22, bold=True), fill=muted)
    draw.text((1518, 82), "5 / 5", font=_font(22, bold=True), fill=muted, anchor="ra")

    def panel(y0: int, y1: int, title: str) -> None:
        draw.rounded_rectangle((70, y0, 1530, y1), radius=13, fill=(252, 253, 250), outline=(145, 153, 154), width=2)
        draw.rectangle((70, y0, 1530, y0 + 44), fill=(226, 234, 236))
        draw.text((92, y0 + 10), title, font=_font(21, bold=True), fill=ink)

    def check(x: int, y: int, label: str, checked: bool, *, radio: bool = False, color=ink) -> None:
        if radio:
            draw.ellipse((x, y, x + 27, y + 27), outline=color, width=3)
            if checked:
                draw.ellipse((x + 7, y + 7, x + 20, y + 20), fill=color)
        else:
            draw.rectangle((x, y, x + 27, y + 27), outline=color, width=3)
            if checked:
                draw.line((x + 5, y + 14, x + 12, y + 22, x + 24, y + 5), fill=color, width=4, joint="curve")
        draw.text((x + 39, y - 1), label, font=_font(22, bold=checked), fill=color)

    panel(205, 445, "A. CONTROLLING DECISION - SELECT ONE STATE PER LINE")
    draw.text((98, 270), "PHYSICAL CUTOVER", font=_font(24, bold=True), fill=ink)
    check(710, 266, "GO", True, color=green)
    check(965, 266, "HOLD", False)
    draw.text((98, 350), "KESTREL CREDENTIAL DECOMMISSION", font=_font(24, bold=True), fill=ink)
    check(710, 346, "RELEASE", False)
    check(965, 346, "HOLD", True, color=red)
    draw.text((1220, 349), "until C-3", font=_font(22, italic=True), fill=red)

    panel(470, 730, "B. INVENTORY INCLUDED IN THIS RELEASE")
    check(100, 535, "Phoenix PX-771 - 48 eligible pallets", True, color=green)
    check(100, 600, "Tucson TZ-219 - 26 eligible pallets", True, color=green)
    check(100, 665, "QA-61 quarantine - six cases", False, color=red)
    draw.text((890, 668), "EXCLUDED / retain segregated custody", font=_font(21, bold=True), fill=red)

    panel(755, 1145, "C. REQUIRED GATES AND ROLLBACK SETTING")
    check(100, 820, "C-1  TV-19 delta <=0.3 C for 30 continuous min before T-04", True)
    check(100, 885, "C-2  QA-61 cases remain segregated before T-05", True)
    check(100, 950, "C-3  EX-07 signed + REC-2048 zero unresolved before T-06", True)
    check(100, 1015, "C-4  Invoke RB-12 at selected continuous deviation", True)
    draw.text((100, 1074), "ROLLBACK THRESHOLD - SELECT ONE", font=_font(20, bold=True), fill=muted)
    check(520, 1070, ">0.5 C / 5 min", False, radio=True)
    check(840, 1070, ">1.0 C / 10 continuous min", True, radio=True, color=red)
    check(1250, 1070, ">1.5 C / 15 min", False, radio=True)

    panel(1170, 1365, "D. EXECUTED START-TIME CORRECTION")
    draw.text((100, 1240), "CUTOVER START:", font=_font(23, bold=True), fill=ink)
    old_text = "14 JUL 2026 21:30 MST"
    old_x = 410
    draw.text((old_x, 1238), old_text, font=_font(24), fill=muted)
    old_bounds = draw.textbbox((old_x, 1238), old_text, font=_font(24))
    draw.line((old_bounds[0] - 4, (old_bounds[1] + old_bounds[3]) // 2, old_bounds[2] + 4, (old_bounds[1] + old_bounds[3]) // 2), fill=red, width=5)
    draw.text((870, 1230), "14 JUL 2026 22:00 MST", font=_font(27, bold=True, italic=True), fill=blue_ink)
    draw.text((1170, 1282), "AK / 10 JUL 09:18", font=_font(22, italic=True), fill=blue_ink)

    panel(1390, 1765, "E. RECORDED SIGNATURES")
    signature_rows = [
        ("EXECUTIVE SPONSOR", "Mara Velez", "10 Jul 2026 09:12 MST", "APPROVED"),
        ("CUTOVER LEAD", "Avery Kim", "10 Jul 2026 09:18 MST", "ACCEPTED"),
        ("QUALITY", "Dara Nwosu", "10 Jul 2026 09:26 MST", "APPROVED - QA-61 EXCLUDED"),
    ]
    for index, (role, name, recorded, decision) in enumerate(signature_rows):
        y = 1460 + index * 92
        draw.text((100, y), role, font=_font(19, bold=True), fill=muted)
        draw.line((365, y + 34, 735, y + 34), fill=(135, 142, 145), width=2)
        draw.text((390, y - 5), name, font=_font(27, bold=True, italic=True), fill=blue_ink)
        draw.text((770, y), recorded, font=_font(20), fill=ink)
        draw.text((1110, y), decision, font=_font(19, bold=True), fill=green if index < 2 else red)

    draw.text((76, 1812), "Effective record: 10 Jul 2026 09:30 MST", font=_font(20, bold=True), fill=ink)
    draw.text((1518, 1812), "CONTROLLED - SV-2048 / E", font=_font(20, bold=True), fill=muted, anchor="ra")
    image.save(path, format="PNG", optimize=True)


def _configure_document(document: Document) -> None:
    _configure_section(document.sections[0])
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.2)
    normal.paragraph_format.space_after = Pt(0)
    properties = document.core_properties
    properties.title = TITLE
    properties.subject = "Controlled supplier transition file SV-2048"
    properties.author = "Northstar Cold Chain, Regional Operations"
    properties.keywords = "SV-2048, supplier transition, cold storage, cutover"
    properties.created = datetime(2026, 7, 8, 16, 20, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 10, 16, 30, tzinfo=timezone.utc)


def _build_docx(path: Path, stamp_path: Path, graph_path: Path, authorization_path: Path) -> None:
    document = Document()
    _configure_document(document)

    # Page 1 - imported presentation cover/precedence sheet. The original table
    # has decomposed into independent, non-row-ordered text boxes.
    _page_title(document, 1, "Controlled precedence sheet", "REV E / CONTROLLED")
    _add_floating_cell(document, "DOCUMENT CONTROL", x_inches=0.52, y_inches=2.08, width_inches=2.1, height_inches=0.3, font_size=7.2, fill="EDF0F1", bold=True)
    _add_spatial_table(
        document,
        CONTROL_HEADERS,
        CONTROL_ROWS,
        [0.8, 2.45, 0.8, 2.75],
        x=0.52,
        y=2.42,
        header_height=0.34,
        row_height=0.29,
        body_height=0.48,
        seed=101,
        damage_sizes={(0, 1): 19.0, (0, 3): 16.5, (1, 1): 21.0, (1, 3): 22.5, (2, 1): 18.5, (2, 3): 17.5},
        cell_offsets={(0, 1): (0.16, 0.04), (0, 3): (-0.22, 0.08), (1, 1): (0.26, -0.08), (1, 3): (-0.14, 0.05), (2, 1): (-0.1, 0.08), (2, 3): (-0.38, -0.07)},
    )
    _add_floating_cell(
        document,
        "RECORD AUTHORITY - HIGHEST APPLICABLE ROW WINS",
        x_inches=0.55,
        y_inches=4.15,
        width_inches=4.25,
        height_inches=0.34,
        font_size=7.4,
        fill="E7EEF2",
        bold=True,
    )
    _add_spatial_table(
        document,
        AUTHORITY_HEADERS,
        AUTHORITY_ROWS,
        [0.48, 1.7, 2.65, 2.55],
        x=0.5,
        y=4.42,
        header_height=0.34,
        row_height=0.35,
        body_height=0.58,
        seed=173,
        omitted_headers={2, 3},
        damage_sizes={
            (0, 1): 15.0,
            (0, 2): 18.5,
            (0, 3): 16.0,
            (1, 1): 17.0,
            (1, 2): 18.0,
            (1, 3): 18.5,
            (2, 1): 16.0,
            (2, 2): 20.0,
            (2, 3): 17.0,
            (3, 1): 15.5,
            (3, 2): 18.0,
            (3, 3): 16.5,
        },
        cell_offsets={(0, 1): (-0.12, 0.06), (0, 2): (0.2, 0.08), (1, 1): (0.18, -0.08), (1, 3): (-0.28, -0.07), (2, 1): (0.18, 0.08), (2, 2): (-0.18, -0.05), (3, 2): (-0.24, -0.09), (3, 3): (0.12, 0.06)},
        bold_columns={0},
    )
    _add_floating_panel(
        document,
        "CONTROLLED PRECEDENCE SHEET",
        x_twips=520,
        y_twips=930,
        width_inches=7.25,
        height_inches=0.62,
        font_size=39,
        fill="EDF0F1",
        color=RGBColor(24, 47, 61),
    )
    _add_floating_panel(
        document,
        "CONTROLS",
        x_twips=6100,
        y_twips=5150,
        width_inches=2.6,
        height_inches=0.47,
        font_size=25,
        fill="D7E5EB",
        color=RGBColor(39, 83, 108),
    )
    _add_floating_panel(
        document,
        "DOES NOT CONTROL",
        x_twips=160,
        y_twips=5750,
        width_inches=3.25,
        height_inches=0.55,
        font_size=22,
        fill="F3DCD5",
        color=RGBColor(143, 58, 48),
    )
    _add_floating_cell(
        document,
        "LOWER-PRIORITY RECORDS MAY ADD DETAIL",
        x_inches=0.82,
        y_inches=6.72,
        width_inches=3.75,
        height_inches=0.42,
        font_size=13.5,
        font_name="Aptos Display Condensed",
        fill="ECE7D7",
        bold=True,
    )
    _add_floating_cell(
        document,
        "BUT CANNOT REVERSE THE HIGHEST APPLICABLE CONTROLLED STATE",
        x_inches=3.35,
        y_inches=6.96,
        width_inches=4.15,
        height_inches=0.42,
        font_size=11.5,
        font_name="Aptos Display Condensed",
        fill="F3DCD5",
        bold=True,
    )
    document.add_page_break()

    # Page 2 - every schedule cell survives as a separate native object. The
    # visible coordinates preserve bindings; object extraction order does not.
    _page_title(document, 2, "Grouped cutover workplan", "CONTROLLED / WORKPLAN")
    _add_floating_cell(document, "MST THROUGHOUT - GATES CONTROL CONTINUED GROUPS", x_inches=0.42, y_inches=1.58, width_inches=4.1, height_inches=0.3, font_size=7.1, fill="EDF0F1", bold=True)
    _add_floating_cell(document, "WORKBOOK RANGE WP-14", x_inches=0.5, y_inches=2.02, width_inches=2.0, height_inches=0.3, font_size=7.0, fill="E7EEF2", bold=True)
    _add_spatial_table(
        document,
        WORKPLAN_HEADERS,
        WORKPLAN_ROWS,
        [0.4, 0.62, 0.92, 1.75, 0.72, 0.9, 1.82, 0.98],
        x=0.17,
        y=2.4,
        header_height=0.34,
        row_height=0.34,
        body_height=0.6,
        seed=227,
        omitted_headers={3, 6},
        damage_sizes={
            (0, 3): 15.0,
            (0, 6): 13.5,
            (1, 2): 12.0,
            (1, 3): 16.5,
            (1, 6): 15.0,
            (2, 3): 18.0,
            (2, 7): 17.0,
            (3, 6): 15.0,
            (4, 3): 16.0,
            (4, 7): 17.5,
            (5, 3): 15.0,
            (5, 6): 18.0,
        },
        cell_offsets={
            (0, 3): (0.18, 0.04),
            (1, 6): (-0.14, -0.06),
            (2, 2): (0.12, -0.04),
            (2, 7): (-0.18, 0.03),
            (3, 3): (-0.1, -0.05),
            (4, 6): (-0.2, 0.04),
            (5, 3): (0.13, -0.05),
        },
        bold_columns={0, 1, 7},
    )
    _add_floating_cell(document, "HANDOFF MATRIX - CURRENT TRIGGERS", x_inches=0.45, y_inches=5.28, width_inches=3.25, height_inches=0.3, font_size=7.0, fill="E7EEF2", bold=True)
    _add_spatial_table(
        document,
        HANDOFF_HEADERS,
        HANDOFF_ROWS,
        [1.3, 1.45, 1.1, 3.1],
        x=0.42,
        y=5.62,
        header_height=0.32,
        row_height=0.27,
        body_height=0.46,
        seed=281,
        damage_sizes={(0, 3): 12.5, (1, 1): 11.0, (2, 3): 14.0},
        cell_offsets={(0, 3): (-0.12, 0.04), (2, 3): (0.16, -0.05)},
        bold_columns={0},
    )
    _add_floating_panel(
        document,
        "ACTIVITY",
        x_twips=3600,
        y_twips=2600,
        width_inches=2.15,
        height_inches=0.58,
        font_size=28,
        fill="D8E7EC",
        color=RGBColor(40, 91, 118),
    )
    _add_floating_panel(
        document,
        "GATE / EVIDENCE",
        x_twips=9200,
        y_twips=3400,
        width_inches=2.8,
        height_inches=0.65,
        font_size=23,
        fill="F1DFCF",
    )
    _add_floating_cell(
        document,
        "ROLLBACK ACTION",
        x_inches=0.4,
        y_inches=7.02,
        width_inches=1.7,
        height_inches=0.34,
        font_size=8.0,
        fill="F2D9D5",
        color=RGBColor(143, 58, 48),
        bold=True,
    )
    _add_floating_cell(
        document,
        "RB-12: STOP THE ACTIVE TRANSFER",
        x_inches=0.6,
        y_inches=7.38,
        width_inches=2.65,
        height_inches=0.46,
        font_size=12.5,
        font_name="Aptos Narrow Variable",
        fill="F4E8CE",
        bold=True,
    )
    _add_floating_cell(
        document,
        "PRESERVE CUSTODY AT THE LAST CONFIRMED LOCATION",
        x_inches=2.75,
        y_inches=7.22,
        width_inches=3.15,
        height_inches=0.48,
        font_size=11.5,
        font_name="Aptos Narrow Variable",
        fill="E7EEF2",
        bold=True,
    )
    _add_floating_cell(
        document,
        "REPORT THE LAST COMPLETED TASK ON BR-2048",
        x_inches=5.2,
        y_inches=7.48,
        width_inches=2.7,
        height_inches=0.5,
        font_size=12.5,
        font_name="Aptos Narrow Variable",
        fill="F3DCD5",
        bold=True,
    )
    document.add_page_break()

    # Page 3 - submitted/current cells and invoice fields are independent text
    # boxes; the validation stamp remains a displaced raster object.
    _page_title(document, 3, "Submitted-to-approved commercial reconciliation", "FINANCE / APPROVED")
    _add_floating_cell(document, "REDLINE RECONCILIATION - STRUCK VALUES ARE SUPERSEDED", x_inches=0.45, y_inches=1.92, width_inches=4.2, height_inches=0.32, font_size=7.0, fill="EDF0F1", bold=True)
    recon_display_rows = [
        [
            row[0],
            "".join(f"{character}\u0336" if not character.isspace() else character for character in row[1]) if index != 2 else row[1],
            row[2],
            row[3],
        ]
        for index, row in enumerate(RECON_ROWS)
    ]
    _add_spatial_table(
        document,
        RECON_HEADERS,
        recon_display_rows,
        [1.25, 1.95, 2.05, 2.25],
        x=0.42,
        y=2.45,
        header_height=0.35,
        row_height=0.39,
        body_height=0.64,
        seed=337,
        omitted_headers={1, 2},
        damage_sizes={(0, 1): 14.0, (0, 2): 16.0, (1, 1): 15.0, (1, 2): 16.0, (3, 1): 14.5, (4, 2): 17.5},
        cell_offsets={(0, 2): (-0.16, 0.04), (1, 1): (0.12, -0.05), (3, 2): (-0.1, -0.04), (4, 1): (0.15, 0.03)},
        bold_columns={0, 2},
    )
    _add_floating_cell(document, "CURRENT INVOICE ROUTE", x_inches=0.45, y_inches=5.12, width_inches=2.4, height_inches=0.3, font_size=7.0, fill="E7EEF2", bold=True)
    _add_spatial_table(
        document,
        INVOICE_HEADERS,
        INVOICE_ROWS,
        [1.45, 1.05, 1.05, 1.45, 2.35],
        x=0.4,
        y=5.46,
        header_height=0.32,
        row_height=0.28,
        body_height=0.48,
        seed=389,
        damage_sizes={(0, 4): 13.0, (1, 0): 11.5, (1, 3): 12.5, (2, 4): 14.0},
        cell_offsets={(0, 4): (-0.14, 0.04), (1, 3): (0.1, -0.05), (2, 4): (-0.18, -0.03)},
        bold_columns={0},
    )
    _add_floating_panel(
        document,
        "APPROVED / CURRENT",
        x_twips=5050,
        y_twips=2450,
        width_inches=3.1,
        height_inches=0.62,
        font_size=27,
        fill="D9EBDD",
        color=RGBColor(38, 103, 70),
    )
    _add_floating_panel(
        document,
        "SUBMITTED / SUPERSEDED",
        x_twips=60,
        y_twips=3900,
        width_inches=3.65,
        height_inches=0.62,
        font_size=24,
        fill="F2D9D5",
        color=RGBColor(145, 58, 48),
    )
    _add_floating_picture(document, stamp_path, x_twips=2100, y_twips=3600, width_inches=6.15)
    _add_floating_cell(
        document,
        "INVOICE ROUTING",
        x_inches=4.72,
        y_inches=5.82,
        width_inches=3.05,
        height_inches=0.68,
        font_size=24.0,
        font_name="Aptos Display Condensed",
        fill="DDE5E8",
        color=RGBColor(45, 75, 91),
        bold=True,
    )

    landscape_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(landscape_section, landscape=True)
    # Page 4 - arrow geometry is raster, connector labels are shuffled native
    # fragments, and the exception register is a decomposed text-box grid.
    _page_title(document, 4, "Dependency graph and exception register", "REV E / CONTROLLED")
    _label(document, "Directed release logic - Rev E control paths")
    # Writer cannot keep the oversized Draw canvas and its detached register
    # objects on the logical slide.  The title remains on physical page 4, the
    # graph spills to page 5, and the register spills again to page 6.
    document.add_picture(str(graph_path), width=Inches(9.85))
    document.paragraphs[-1].paragraph_format.space_after = Pt(0)
    _add_edge_label_fragments(document)
    _add_floating_cell(document, "CONTROL PATH LABEL REGISTER - REV E", x_inches=1.42, y_inches=5.0, width_inches=6.15, height_inches=0.46, font_size=21.0, font_name="Aptos Display Condensed", fill="C9D9E1", color=RGBColor(44, 79, 96), bold=True)
    _add_spatial_table(
        document,
        ["Exception", "Affected obligation", "Current finding", "Owner", "Due", "", ""],
        EXCEPTION_ROWS,
        [0.65, 1.25, 2.45, 0.9, 1.15, 2.6, 1.05],
        x=0.32,
        y=6.03,
        header_height=0.31,
        row_height=0.19,
        body_height=0.36,
        seed=443,
        damage_sizes={
            (0, 2): 10.8,
            (0, 5): 12.5,
            (1, 2): 11.5,
            (1, 5): 13.5,
            (2, 2): 12.5,
            (2, 5): 12.0,
            (3, 2): 11.0,
            (4, 2): 12.0,
            (4, 5): 13.0,
            (5, 2): 11.5,
            (5, 5): 12.5,
        },
        cell_offsets={(0, 5): (-0.14, 0.03), (1, 2): (0.12, -0.04), (2, 5): (-0.18, 0.04), (4, 2): (0.1, -0.05), (5, 5): (-0.12, -0.03)},
        bold_columns={0, 6},
    )
    _add_floating_panel(document, "ACCEPTED CLOSURE EVIDENCE", x_twips=7900, y_twips=7900, width_inches=3.7, height_inches=0.48, font_size=22, fill="F4E8CE")
    _add_floating_panel(document, "STATE", x_twips=13500, y_twips=8300, width_inches=1.45, height_inches=0.52, font_size=27, fill="F3DCD5", color=RGBColor(143, 58, 48))

    final_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(final_section)
    # Logical page 5 intentionally remains a clean 220-DPI executed scan on
    # physical page 7. It controls the damaged lower-priority imported pages.
    _page_title(document, 5, "Executed selective authorization", "EFFECTIVE")
    document.add_picture(str(authorization_path), width=Inches(7.28))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(path)


def _export_with_libreoffice(docx_path: Path, pdf_path: Path, work_dir: Path) -> None:
    soffice = shutil.which("soffice") or "/opt/homebrew/bin/soffice"
    if not Path(soffice).exists():
        raise RuntimeError("LibreOffice soffice is required to generate P23")
    export_dir = work_dir / "export"
    profile_dir = work_dir / "lo-profile"
    export_dir.mkdir(parents=True, exist_ok=True)
    profile_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    env.update({"LC_ALL": "C", "LANG": "C", "TZ": "UTC", "SOURCE_DATE_EPOCH": "1783696200"})
    result = subprocess.run(
        [
            soffice,
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(export_dir),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=90,
        env=env,
    )
    converted = export_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not converted.exists():
        raise RuntimeError(
            "LibreOffice PDF export failed: "
            + (result.stderr.strip() or result.stdout.strip() or f"exit {result.returncode}")
        )

    normalized = work_dir / "metadata-normalized.pdf"
    reader = PdfReader(str(converted))
    if len(reader.pages) != 7:
        raise ValueError(f"{CASE_ID}: LibreOffice generated {len(reader.pages)} pages; expected 7")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": TITLE,
            "/Author": "Northstar Cold Chain, Regional Operations",
            "/Subject": "Controlled supplier transition file SV-2048",
            "/Creator": "LibreOffice Writer",
            "/Producer": "LibreOffice PDF export",
        }
    )
    with normalized.open("wb") as stream:
        writer.write(stream)

    qpdf = shutil.which("qpdf")
    if not qpdf:
        raise RuntimeError("qpdf is required to normalize P23 deterministically")
    subprocess.run(
        [qpdf, "--deterministic-id", "--object-streams=generate", str(normalized), str(pdf_path)],
        check=True,
        capture_output=True,
        text=True,
    )


GOLD_SECTIONS = {
    1: "1. Controlled precedence sheet",
    2: "2. Grouped cutover workplan",
    3: "3. Submitted-to-approved commercial reconciliation",
    4: "4. Dependency graph and exception register",
    5: "4. Dependency graph and exception register",
    6: "4. Dependency graph and exception register",
    7: "5. Executed selective authorization",
}


def _source_anchor(page: int, label: str, layer: str) -> dict:
    document_page = 4 if page in {4, 5, 6} else 5 if page == 7 else page
    return {
        "page": page,
        "documentPage": document_page,
        "layer": layer,
        "sectionPath": [TITLE, GOLD_SECTIONS[page], label],
    }


def _closed_table(keys: list[str]) -> dict:
    return {"scope": "table_rows", "keys": keys}


def _region(
    region_id: str,
    label: str,
    kind: str,
    leaves: list[dict],
    *,
    primary_axis: str,
    page: int | None = None,
    layer: str = "native_text",
    modality: str | None = None,
    secondary_axes: tuple[str, ...] = (),
    text_only_recoverable: bool = True,
    budget: int = 1,
    unique_evidence: bool = True,
    closed_world: dict | None = None,
    source_anchors: list[dict] | None = None,
    gold_section: str | None = None,
) -> dict:
    if source_anchors is None:
        if page is None:
            raise ValueError(f"{region_id}: page or source_anchors is required")
        source_anchors = [_source_anchor(page, label, layer)]
    primary_page = int(source_anchors[0]["page"])
    region = {
        "id": region_id,
        "label": label,
        "sourceAnchors": source_anchors,
        "goldSection": gold_section or GOLD_SECTIONS[primary_page],
        "kind": kind,
        "modality": modality or layer,
        "uniqueEvidence": unique_evidence,
        "primaryAxis": primary_axis,
        "secondaryAxes": list(secondary_axes),
        "textOnlyRecoverable": text_only_recoverable,
        "budget": budget,
        "leaves": leaves,
    }
    if closed_world is not None:
        region["closedWorld"] = closed_world
    return region



def _regions() -> list[dict]:
    authority_bindings = {
        (row[1], column)
        for row in AUTHORITY_ROWS
        for column in ("Controls", "Does not control")
    }
    workplan_bindings = {
        ("T-01", "Activity"),
        ("T-01", "Gate / evidence"),
        ("T-02", "Activity"),
        ("T-02", "Gate / evidence"),
        ("T-02", "State"),
        ("T-04", "Activity"),
        ("T-04", "Gate / evidence"),
        ("T-04", "State"),
        ("T-03", "Gate / evidence"),
        ("T-05", "Activity"),
        ("T-05", "Gate / evidence"),
        ("T-05", "State"),
        ("T-06", "Activity"),
        ("T-06", "Gate / evidence"),
        ("T-06", "State"),
    }
    handoff_bindings = {(row[0], "Required content") for row in HANDOFF_ROWS}
    recon_bindings = {(row[0], "Approved / current") for row in RECON_ROWS} | {
        ("Boreal implementation", "Submitted / superseded"),
        ("Kestrel overlap", "Submitted / superseded"),
        ("Contingency", "Submitted / superseded"),
        ("Request / ceiling", "Submitted / superseded"),
        ("Kestrel overlap", "Approval note"),
        ("Contingency", "Approval note"),
        ("Request / ceiling", "Approval note"),
    }
    invoice_bindings = {(row[0], "PO") for row in INVOICE_ROWS} | {
        ("Boreal Storage", "Route"),
        ("Kestrel Logistics", "Route"),
        ("ThermoVision", "Cost center"),
        ("ThermoVision", "Route"),
    }
    exception_bindings = {(row[0], "State") for row in EXCEPTION_ROWS} | {
        ("E-05", "Accepted closure evidence"),
        ("EX-07", "Affected obligation"),
        ("EX-07", "Accepted closure evidence"),
        ("E-08", "Current finding"),
        ("E-08", "Accepted closure evidence"),
        ("E-11", "Current finding"),
        ("E-14", "Current finding"),
    }
    return [
        _region(
            "p01.control",
            "Damaged document-control grid",
            "table",
            [
                leaf("p01.control.identity", "The controlled file is SV-2048 at Revision E.", evidence=["SV-2048", "Revision E"]),
                leaf("p01.control.issued", "SV-2048 was issued 12 Jul 2026 at 18:40 MST.", evidence=["SV-2048", "12 Jul 2026", "18:40 MST"]),
                leaf("p01.control.start", "The controlled cutover start is 14 Jul 2026 at 22:00 MST.", evidence=["14 Jul 2026", "22:00 MST"]),
                leaf("p01.control.scope", "The controlled lots are PX-771 and TZ-219, coordinated by Avery Kim.", evidence=["PX-771", "TZ-219", "Avery Kim"]),
            ],
            page=1,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="native_layer_recovery",
            secondary_axes=("table_reconstruction", "precise_recall"),
            text_only_recoverable=False,
            budget=1,
        ),
        _region(
            "p01.authority",
            "Overlapped authority table and detached headings",
            "table",
            [
                *table_leaves(
                    "p01.authority",
                    AUTHORITY_HEADERS,
                    AUTHORITY_ROWS,
                    key_column=1,
                    consequential={(row[1], "Controls") for row in AUTHORITY_ROWS},
                    scored_bindings=authority_bindings,
                ),
                leaf(
                    "p01.authority.rule",
                    "A lower-priority record may add detail but cannot reverse the highest applicable controlled state.",
                    harm=2,
                    evidence=["lower-priority", "detail", "cannot reverse", "highest applicable"],
                ),
            ],
            page=1,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="source_precedence",
            secondary_axes=("native_layer_recovery", "table_reconstruction", "reading_order"),
            text_only_recoverable=False,
            budget=3,
            closed_world=_closed_table([row[1] for row in AUTHORITY_ROWS]),
        ),
        _region(
            "p02.workplan",
            "Two-tier clipped workplan with continued groups",
            "table",
            table_leaves(
                "p02.workplan",
                WORKPLAN_HEADERS,
                WORKPLAN_ROWS,
                consequential={
                    ("T-02", "Gate / evidence"),
                    ("T-04", "State"),
                    ("T-05", "Activity"),
                    ("T-05", "State"),
                    ("T-06", "Gate / evidence"),
                    ("T-06", "State"),
                },
                scored_bindings=workplan_bindings,
            ),
            page=2,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="native_layer_recovery",
            secondary_axes=("table_reconstruction", "precise_recall", "reading_order"),
            text_only_recoverable=False,
            budget=4,
            closed_world=_closed_table([row[0] for row in WORKPLAN_ROWS]),
        ),
        _region(
            "p02.handoff",
            "Clipped handoff matrix",
            "table",
            table_leaves(
                "p02.handoff",
                HANDOFF_HEADERS,
                HANDOFF_ROWS,
                consequential={("RB-12 invoked", "Required content")},
                scored_bindings=handoff_bindings,
            ),
            page=2,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="table_reconstruction",
            secondary_axes=("native_layer_recovery", "precise_recall"),
            text_only_recoverable=False,
            budget=2,
            closed_world=_closed_table([row[0] for row in HANDOFF_ROWS]),
        ),
        _region(
            "p02.rollback",
            "Rollback action",
            "text",
            [
                leaf("p02.rollback.stop", "Invoking RB-12 stops the active transfer.", harm=2),
                leaf("p02.rollback.custody", "RB-12 preserves custody at the last confirmed location.", harm=2),
                leaf("p02.rollback.report", "RB-12 requires the last completed task to be reported on BR-2048."),
            ],
            page=2,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="native_layer_recovery",
            secondary_axes=("precise_recall", "long_context_coherence"),
            text_only_recoverable=False,
            budget=2,
        ),
        _region(
            "p03.reconciliation",
            "Overlapped submitted-to-current reconciliation",
            "table",
            [
                *table_leaves(
                    "p03.reconciliation",
                    RECON_HEADERS,
                    RECON_ROWS,
                    consequential={
                        ("Boreal implementation", "Approved / current"),
                        ("Kestrel overlap", "Approved / current"),
                        ("Contingency", "Approved / current"),
                        ("Request / ceiling", "Approved / current"),
                    },
                    scored_bindings=recon_bindings,
                ),
                form_state_leaf("p03.reconciliation.boreal-old", "Q-884 rev B / $19,200 is crossed out.", "Q-884 rev B / $19,200", "crossed"),
                form_state_leaf("p03.reconciliation.overlap-old", "3 nights / $6,480 is crossed out.", "3 nights / $6,480", "crossed"),
                form_state_leaf("p03.reconciliation.contingency-old", "$4,500 contingency is crossed out.", ["$4,500 contingency", "$4,500"], "crossed"),
                form_state_leaf("p03.reconciliation.total-old", "Submitted total / $32,330 is crossed out.", ["Submitted total / $32,330", "$32,330"], "crossed"),
            ],
            page=3,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="native_layer_recovery",
            secondary_axes=("table_reconstruction", "form_state", "precise_recall"),
            text_only_recoverable=False,
            budget=4,
            closed_world=_closed_table([row[0] for row in RECON_ROWS]),
        ),
        _region(
            "p03.invoice",
            "Displaced invoice-route table",
            "table",
            table_leaves("p03.invoice", INVOICE_HEADERS, INVOICE_ROWS, scored_bindings=invoice_bindings),
            page=3,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="table_reconstruction",
            secondary_axes=("native_layer_recovery", "precise_recall"),
            text_only_recoverable=False,
            budget=2,
            closed_world=_closed_table([row[0] for row in INVOICE_ROWS]),
        ),
        _region(
            "p03.stamp",
            "Displaced raster finance stamp",
            "image",
            [
                visual_leaf("p03.stamp.signer", "The validation stamp records Jonah Mercer on 09 Jul 2026 at 16:42 MST.", [["Jonah Mercer"], ["09 Jul 2026"], ["16:42 MST"]]),
                visual_leaf("p03.stamp.control", "The validation stamp identifies approved reconciliation control RC-2048-E.", [["approved reconciliation"], ["RC-2048-E"]]),
                visual_leaf("p03.stamp.boundary", "The validation stamp says cost validation is not operational release.", [["cost validation"], ["not operational release"]], harm=2),
            ],
            page=3,
            layer="raster",
            modality="raster",
            primary_axis="image_description",
            secondary_axes=("mixed_modality_fusion", "precise_recall"),
            text_only_recoverable=False,
            budget=2,
        ),
        _region(
            "p04.graph-operational",
            "Directed operational paths under broken object stacking",
            "diagram",
            [
                directed_edge_leaf(
                    f"p04.graph.edge-{index:02d}",
                    f"A directed {relation} edge runs from {source} to {destination}.",
                    [source],
                    [destination],
                    relation=[relation],
                    identifier=[f"E{index:02d}"],
                    harm=2 if relation in {"gates", "releases", "eligible only", "waiver", "close access", "breach", "halts"} else 1,
                )
                for index, (source, destination, relation) in enumerate(GRAPH_EDGES[:7], start=1)
            ],
            source_anchors=[
                _source_anchor(4, "Orphaned logical-page title", "native_layer_recovery"),
                _source_anchor(5, "Displaced directed graph objects", "raster"),
                _source_anchor(5, "Detached native connector-label fragments A", "native_layer_recovery"),
                _source_anchor(6, "Detached native connector-label fragments B", "native_layer_recovery"),
            ],
            modality="mixed",
            primary_axis="mixed_modality_fusion",
            secondary_axes=("image_description", "structure_reconstruction", "precise_recall"),
            unique_evidence=False,
            text_only_recoverable=False,
            budget=4,
        ),
        _region(
            "p04.graph-control",
            "Directed closeout and rollback paths under broken object stacking",
            "diagram",
            [
                directed_edge_leaf(
                    f"p04.graph.edge-{index:02d}",
                    f"A directed {relation} edge runs from {source} to {destination}.",
                    [source],
                    [destination],
                    relation=[relation],
                    identifier=[f"E{index:02d}"],
                    harm=2 if relation in {"gates", "releases", "eligible only", "waiver", "close access", "breach", "halts"} else 1,
                )
                for index, (source, destination, relation) in enumerate(GRAPH_EDGES[7:], start=8)
            ],
            source_anchors=[
                _source_anchor(4, "Orphaned logical-page title", "native_layer_recovery"),
                _source_anchor(5, "Displaced directed graph objects", "raster"),
                _source_anchor(5, "Detached native connector-label fragments A", "native_layer_recovery"),
                _source_anchor(6, "Detached native connector-label fragments B", "native_layer_recovery"),
            ],
            modality="mixed",
            primary_axis="mixed_modality_fusion",
            secondary_axes=("image_description", "structure_reconstruction", "precise_recall"),
            unique_evidence=False,
            text_only_recoverable=False,
            budget=4,
        ),
        _region(
            "p04.exceptions",
            "Clipped exception register under detached headings",
            "table",
            table_leaves(
                "p04.exceptions",
                EXCEPTION_HEADERS,
                EXCEPTION_ROWS,
                consequential={
                    ("EX-07", "Affected obligation"),
                    ("EX-07", "Accepted closure evidence"),
                    ("EX-07", "State"),
                    ("E-08", "Current finding"),
                    ("E-08", "State"),
                },
                scored_bindings=exception_bindings,
            ),
            page=6,
            layer="native_layer_recovery",
            modality="native_layer_recovery",
            primary_axis="native_layer_recovery",
            secondary_axes=("table_reconstruction", "source_precedence", "precise_recall"),
            text_only_recoverable=False,
            budget=4,
            closed_world=_closed_table([row[0] for row in EXCEPTION_ROWS]),
        ),
        _region(
            "p05.decision",
            "Executed selective decision states",
            "form",
            [
                form_state_leaf("p05.decision.cutover-go", "Physical cutover - GO is checked.", ["Physical cutover - GO", "Physical cutover GO"], "checked", harm=2),
                form_state_leaf("p05.decision.cutover-hold", "Physical cutover - HOLD is unchecked.", ["Physical cutover - HOLD", "Physical cutover HOLD"], "unchecked"),
                form_state_leaf("p05.decision.credentials-release", "Kestrel credential decommission - RELEASE is unchecked.", ["Kestrel credential decommission - RELEASE", "credential decommission RELEASE"], "unchecked"),
                form_state_leaf("p05.decision.credentials-hold", "Kestrel credential decommission - HOLD until C-3 is checked.", "Kestrel credential decommission - HOLD until C-3", "checked", harm=2),
            ],
            page=7,
            layer="raster",
            modality="raster",
            primary_axis="source_precedence",
            secondary_axes=("form_state", "image_description"),
            text_only_recoverable=False,
            budget=3,
        ),
        _region(
            "p05.inventory",
            "Executed inventory selection",
            "form",
            [
                form_state_leaf("p05.inventory.phx", "Phoenix PX-771 - 48 eligible pallets is checked.", "Phoenix PX-771 - 48 eligible pallets", "checked"),
                form_state_leaf("p05.inventory.tus", "Tucson TZ-219 - 26 eligible pallets is checked.", "Tucson TZ-219 - 26 eligible pallets", "checked"),
                form_state_leaf("p05.inventory.qa61", "QA-61 quarantine - six cases is unchecked and excluded.", "QA-61 quarantine - six cases", "unchecked", harm=2),
            ],
            page=7,
            layer="raster",
            modality="raster",
            primary_axis="image_description",
            secondary_axes=("form_state", "source_precedence"),
            text_only_recoverable=False,
            budget=2,
        ),
        _region(
            "p05.gates",
            "Executed required-gate states",
            "form",
            [
                form_state_leaf("p05.gates.c1", "C-1 TV-19 delta <=0.3 C for 30 continuous min before T-04 is checked.", "C-1 TV-19 delta <=0.3 C for 30 continuous min before T-04", "checked", harm=2),
                form_state_leaf("p05.gates.c2", "C-2 QA-61 cases remain segregated before T-05 is checked.", "C-2 QA-61 cases remain segregated before T-05", "checked", harm=2),
                form_state_leaf("p05.gates.c3", "C-3 EX-07 signed plus REC-2048 zero unresolved before T-06 is checked.", "C-3 EX-07 signed + REC-2048 zero unresolved before T-06", "checked", harm=2),
                form_state_leaf("p05.gates.c4", "C-4 invoke RB-12 at the selected continuous deviation is checked.", "C-4 Invoke RB-12 at selected continuous deviation", "checked"),
            ],
            page=7,
            layer="raster",
            modality="raster",
            primary_axis="image_description",
            secondary_axes=("form_state", "precise_recall"),
            text_only_recoverable=False,
            budget=3,
        ),
        _region(
            "p05.threshold",
            "Selected rollback threshold",
            "form",
            [
                form_state_leaf("p05.threshold.low", ">0.5 C / 5 min is unchecked.", ">0.5 C / 5 min", "unchecked"),
                form_state_leaf("p05.threshold.current", ">1.0 C / 10 continuous min is checked.", ">1.0 C / 10 continuous min", "checked", harm=2),
                form_state_leaf("p05.threshold.high", ">1.5 C / 15 min is unchecked.", ">1.5 C / 15 min", "unchecked"),
            ],
            page=7,
            layer="raster",
            modality="raster",
            primary_axis="image_description",
            secondary_axes=("form_state", "precise_recall"),
            text_only_recoverable=False,
            budget=2,
        ),
        _region(
            "p05.correction",
            "Handwritten start-time correction",
            "form",
            [
                form_state_leaf("p05.correction.old", "Cutover start 14 Jul 2026 21:30 MST is crossed out.", ["Cutover start 14 Jul 2026 21:30 MST", "14 Jul 2026 21:30 MST"], "crossed", harm=2),
                visual_leaf("p05.correction.current", "Avery Kim corrected the cutover start to 14 Jul 2026 at 22:00 MST and initialed AK on 10 Jul at 09:18.", [["corrected"], ["14 Jul 2026", "14 JUL 2026"], ["22:00 MST"], ["AK"], ["10 Jul", "10 JUL"], ["09:18"]], harm=2),
            ],
            page=7,
            layer="raster",
            modality="raster",
            primary_axis="image_description",
            secondary_axes=("form_state", "precise_recall"),
            text_only_recoverable=False,
            budget=1,
        ),
        _region(
            "p05.signatures",
            "Executed signatures and effective time",
            "form",
            [
                visual_leaf("p05.signatures.sponsor", "Mara Velez approved as executive sponsor on 10 Jul 2026 at 09:12 MST.", [["Mara Velez"], ["executive sponsor"], ["approved"], ["10 Jul 2026"], ["09:12 MST"]]),
                visual_leaf("p05.signatures.lead", "Avery Kim accepted as cutover lead on 10 Jul 2026 at 09:18 MST.", [["Avery Kim"], ["cutover lead"], ["accepted"], ["10 Jul 2026"], ["09:18 MST"]]),
                visual_leaf("p05.signatures.quality", "Dara Nwosu approved as quality on 10 Jul 2026 at 09:26 MST with QA-61 excluded.", [["Dara Nwosu"], ["quality"], ["approved"], ["10 Jul 2026"], ["09:26 MST"], ["QA-61 excluded"]], harm=2),
                visual_leaf("p05.signatures.effective", "The effective record time is 10 Jul 2026 at 09:30 MST.", [["effective record"], ["10 Jul 2026"], ["09:30 MST"]]),
            ],
            page=7,
            layer="raster",
            modality="raster",
            primary_axis="image_description",
            secondary_axes=("precise_recall", "source_precedence"),
            text_only_recoverable=False,
            budget=2,
        ),
    ]


def _gold() -> str:
    recon_gold_rows = [
        [row[0], f"~~{row[1]}~~" if index != 2 else row[1], row[2], row[3]]
        for index, row in enumerate(RECON_ROWS)
    ]
    edge_lines = "\n".join(f"- `{source} -> {destination}` - {relation}" for source, destination, relation in GRAPH_EDGES)
    return (
        f"# {TITLE}\n\n"
        "## 1. Controlled precedence sheet\n\n"
        + markdown_table(CONTROL_HEADERS, CONTROL_ROWS)
        + "\n\n"
        + markdown_table(AUTHORITY_HEADERS, AUTHORITY_ROWS)
        + "\n\nA lower-priority record may add detail but cannot reverse the highest applicable controlled state.\n\n"
        "## 2. Grouped cutover workplan\n\n"
        + markdown_table(WORKPLAN_HEADERS, WORKPLAN_ROWS)
        + "\n\n"
        + markdown_table(HANDOFF_HEADERS, HANDOFF_ROWS)
        + "\n\nWhen RB-12 is invoked, stop the active transfer, preserve custody at the last confirmed location, and report the last completed task on BR-2048.\n\n"
        "## 3. Submitted-to-approved commercial reconciliation\n\n"
        + markdown_table(RECON_HEADERS, recon_gold_rows)
        + "\n\n"
        + markdown_table(INVOICE_HEADERS, INVOICE_ROWS)
        + "\n\n"
        "Raster finance stamp: Jonah Mercer; 09 Jul 2026 16:42 MST; approved reconciliation; control RC-2048-E; cost validation is not operational release.\n\n"
        "## 4. Dependency graph and exception register\n\n"
        + edge_lines
        + "\n\n"
        + markdown_table(EXCEPTION_HEADERS, EXCEPTION_ROWS)
        + "\n\n## 5. Executed selective authorization\n\n"
        "### Controlling decision\n\n"
        "- [x] Physical cutover - GO\n"
        "- [ ] Physical cutover - HOLD\n"
        "- [ ] Kestrel credential decommission - RELEASE\n"
        "- [x] Kestrel credential decommission - HOLD until C-3\n\n"
        "### Inventory included\n\n"
        "- [x] Phoenix PX-771 - 48 eligible pallets\n"
        "- [x] Tucson TZ-219 - 26 eligible pallets\n"
        "- [ ] QA-61 quarantine - six cases; excluded and retained in segregated custody\n\n"
        "### Required gates\n\n"
        "- [x] C-1 TV-19 delta <=0.3 C for 30 continuous min before T-04\n"
        "- [x] C-2 QA-61 cases remain segregated before T-05\n"
        "- [x] C-3 EX-07 signed + REC-2048 zero unresolved before T-06\n"
        "- [x] C-4 Invoke RB-12 at selected continuous deviation\n\n"
        "Rollback threshold: ( ) >0.5 C / 5 min; (x) >1.0 C / 10 continuous min; ( ) >1.5 C / 15 min.\n\n"
        "### Executed correction\n\n"
        "Cutover start: ~~14 JUL 2026 21:30 MST~~. Avery Kim corrected the handwritten current value to **14 JUL 2026 22:00 MST**, initialed **AK / 10 JUL 09:18**.\n\n"
        "### Recorded signatures\n\n"
        "- Executive sponsor Mara Velez - APPROVED - 10 Jul 2026 09:12 MST\n"
        "- Cutover lead Avery Kim - ACCEPTED - 10 Jul 2026 09:18 MST\n"
        "- Quality Dara Nwosu - APPROVED, QA-61 EXCLUDED - 10 Jul 2026 09:26 MST\n\n"
        "Effective record: 10 Jul 2026 09:30 MST.\n"
    )


def build(output_root):
    output_root = Path(output_root).resolve()
    case_dir = output_root / "cases" / CASE_ID
    case_dir.mkdir(parents=True, exist_ok=True)
    source_pdf = case_dir / "source.pdf"
    with tempfile.TemporaryDirectory(prefix="p23-office-export-") as temp_name:
        work_dir = Path(temp_name)
        docx_path = work_dir / "northstar-supplier-transition-sv-2048.docx"
        stamp_path = work_dir / "finance-validation.png"
        graph_path = work_dir / "dependency-graph.png"
        authorization_path = work_dir / "executed-authorization.png"
        _finance_stamp(stamp_path)
        _dependency_graph(graph_path)
        _executed_authorization(authorization_path)
        _build_docx(docx_path, stamp_path, graph_path, authorization_path)
        _export_with_libreoffice(docx_path, source_pdf, work_dir)

    regions = _regions()
    gold = _gold()
    finalize_fact_regions(regions, gold)
    (case_dir / "gold.md").write_text(gold, encoding="utf-8")
    (case_dir / "facts.json").write_text(
        json.dumps(
            {
                "schemaVersion": 3,
                "id": CASE_ID,
                "title": TITLE,
                "family": FAMILY,
                "tags": TAGS,
                "regions": regions,
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    (case_dir / "spec.md").write_text(
        f"# {TITLE}\n\n"
        "Source modality: five logical presentation pages exported through LibreOffice Writer as a seven-physical-page mixed native/raster PDF. Logical page 4 catastrophically spills: its title remains on physical page 4, the Draw graph lands on physical page 5, and the detached exception register lands on physical page 6. The clean 220-DPI executed form still prints logical 5/5 but appears on physical page 7.\n\n"
        f"Family: `{FAMILY}`\n\n"
        "Purpose: recover the intended logical packet after a realistic Slides/PowerPoint-style office round trip substitutes missing condensed fonts, inflates runs, decomposes tables into independently positioned cells, detaches headers, changes object order, overlaps native text, breaks logical pagination, and collides graphic groups with register content across several pages. The native PDF object stream is intentionally shuffled rather than an intact row-wise table export; visible coordinates and row IDs preserve the bindings. On the dependency spread, native connector-ID/label fragments split across physical pages 5-6 must be joined to visible raster arrow endpoints and direction on physical page 5, while the associated exception register has also spilled to physical page 6, so neither modality nor page is a complete answer key. Physical page 7 remains the clean executed source of truth.\n\n"
        "Every scored fact is directly present in a visible or native PDF object. No answer-key-only value, hidden text, arithmetic inference, tiny-font needle, blur, or arbitrary image degradation is used. The gold document reconstructs bindings, reading order, directed edges, superseded/current values, and executed states from those objects.\n\n"
        "The source PDF contains the office export's own text and image objects. Container normalization changes metadata and serialization only; it does not add, hide, or replace page content.\n\n"
        "Tags: "
        + ", ".join(f"`{tag}`" for tag in TAGS)
        + "\n",
        encoding="utf-8",
    )

    prefix = output_root.relative_to(REPO_ROOT).as_posix()
    base = f"{prefix}/cases/{CASE_ID}"
    return {
        "id": CASE_ID,
        "title": TITLE,
        "family": FAMILY,
        "tags": TAGS,
        "pages": 7,
        "pdf": f"{base}/source.pdf",
        "gold": f"{base}/gold.md",
        "spec": f"{base}/spec.md",
        "facts": f"{base}/facts.json",
    }
